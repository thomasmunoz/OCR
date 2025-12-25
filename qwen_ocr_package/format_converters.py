#!/usr/bin/env python3
"""
Format Converters for Multiple Document Types

Converts various document formats to text for OCR processing.
Supports: DOCX, DOC, XLSX, XLS, TXT, HTML, and more.
"""

from pathlib import Path
from typing import Dict, Any, Optional, List
import io
import tempfile
import subprocess


class FormatConverter:
    """Base class for format converters"""

    @staticmethod
    def convert(file_path: Path) -> str:
        """Convert file to text"""
        raise NotImplementedError


class TextConverter(FormatConverter):
    """Converter for plain text files"""

    @staticmethod
    def convert(file_path: Path) -> str:
        """Read plain text file"""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue

        raise ValueError(f"Could not decode text file: {file_path}")


class DOCXConverter(FormatConverter):
    """Converter for DOCX files"""

    @staticmethod
    def convert(file_path: Path) -> str:
        """Extract text from DOCX"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx not installed. Install with: pip install python-docx"
            )

        doc = Document(file_path)
        text_parts = []

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n".join(text_parts)


class DOCConverter(FormatConverter):
    """Converter for legacy DOC files"""

    @staticmethod
    def convert(file_path: Path) -> str:
        """
        Extract text from DOC files
        Tries multiple methods: antiword, textutil (macOS), or conversion to DOCX
        """
        # Try antiword (Linux/macOS)
        try:
            result = subprocess.run(
                ['antiword', str(file_path)],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

        # Try textutil (macOS only)
        try:
            with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as tmp:
                tmp_path = tmp.name

            subprocess.run(
                ['textutil', '-convert', 'txt', '-output', tmp_path, str(file_path)],
                capture_output=True,
                timeout=30
            )

            with open(tmp_path, 'r') as f:
                text = f.read()

            Path(tmp_path).unlink()
            return text
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

        # Try converting DOC to DOCX using libreoffice
        try:
            with tempfile.TemporaryDirectory() as tmp_dir:
                subprocess.run(
                    ['libreoffice', '--headless', '--convert-to', 'docx',
                     '--outdir', tmp_dir, str(file_path)],
                    capture_output=True,
                    timeout=60
                )

                docx_path = Path(tmp_dir) / f"{file_path.stem}.docx"
                if docx_path.exists():
                    return DOCXConverter.convert(docx_path)
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

        raise ValueError(
            "Could not convert DOC file. Please install one of: "
            "antiword, LibreOffice, or convert to DOCX manually"
        )


class XLSXConverter(FormatConverter):
    """Converter for XLSX files"""

    @staticmethod
    def convert(file_path: Path) -> str:
        """Extract text from XLSX"""
        try:
            from openpyxl import load_workbook
        except ImportError:
            raise ImportError(
                "openpyxl not installed. Install with: pip install openpyxl"
            )

        workbook = load_workbook(file_path, data_only=True)
        text_parts = []

        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text_parts.append(f"\n=== Sheet: {sheet_name} ===\n")

            for row in sheet.iter_rows(values_only=True):
                # Filter out None values and convert to strings
                row_values = [str(cell) for cell in row if cell is not None]
                if row_values:
                    text_parts.append(" | ".join(row_values))

        return "\n".join(text_parts)


class XLSConverter(FormatConverter):
    """Converter for legacy XLS files"""

    @staticmethod
    def convert(file_path: Path) -> str:
        """Extract text from XLS"""
        try:
            import xlrd
        except ImportError:
            raise ImportError(
                "xlrd not installed. Install with: pip install xlrd"
            )

        workbook = xlrd.open_workbook(file_path)
        text_parts = []

        for sheet in workbook.sheets():
            text_parts.append(f"\n=== Sheet: {sheet.name} ===\n")

            for row_idx in range(sheet.nrows):
                row = sheet.row(row_idx)
                row_values = [str(cell.value) for cell in row if cell.value]
                if row_values:
                    text_parts.append(" | ".join(row_values))

        return "\n".join(text_parts)


class HTMLConverter(FormatConverter):
    """Converter for HTML/XHTML files"""

    @staticmethod
    def convert(file_path: Path) -> str:
        """Extract text from HTML"""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            # Fallback to regex-based extraction
            return HTMLConverter._convert_regex(file_path)

        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            html_content = f.read()

        soup = BeautifulSoup(html_content, 'html.parser')

        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()

        # Get text
        text = soup.get_text()

        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)

        return text

    @staticmethod
    def _convert_regex(file_path: Path) -> str:
        """Fallback regex-based HTML text extraction"""
        import re

        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            html_content = f.read()

        # Remove scripts and styles
        text = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL)
        text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL)

        # Remove HTML tags
        text = re.sub(r'<[^>]+>', ' ', text)

        # Clean up whitespace
        text = re.sub(r'\s+', ' ', text)

        return text.strip()


class RTFConverter(FormatConverter):
    """Converter for RTF files"""

    @staticmethod
    def convert(file_path: Path) -> str:
        """Extract text from RTF"""
        try:
            from striprtf.striprtf import rtf_to_text
        except ImportError:
            raise ImportError(
                "striprtf not installed. Install with: pip install striprtf"
            )

        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            rtf_content = f.read()

        return rtf_to_text(rtf_content)


class CSVConverter(FormatConverter):
    """Converter for CSV files"""

    @staticmethod
    def convert(file_path: Path) -> str:
        """Extract text from CSV"""
        import csv

        text_parts = []
        encodings = ['utf-8', 'latin-1', 'cp1252']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding, newline='') as f:
                    reader = csv.reader(f)
                    for row in reader:
                        row_values = [str(cell) for cell in row if cell]
                        if row_values:
                            text_parts.append(" | ".join(row_values))
                break
            except UnicodeDecodeError:
                continue

        if not text_parts:
            raise ValueError(f"Could not decode CSV file: {file_path}")

        return "\n".join(text_parts)


class ConverterRegistry:
    """Registry of format converters"""

    CONVERTERS = {
        '.txt': TextConverter,
        '.text': TextConverter,
        '.docx': DOCXConverter,
        '.doc': DOCConverter,
        '.xlsx': XLSXConverter,
        '.xls': XLSConverter,
        '.html': HTMLConverter,
        '.htm': HTMLConverter,
        '.xhtml': HTMLConverter,
        '.rtf': RTFConverter,
        '.csv': CSVConverter,
    }

    @classmethod
    def get_converter(cls, file_path: Path) -> Optional[FormatConverter]:
        """Get appropriate converter for file"""
        ext = file_path.suffix.lower()
        return cls.CONVERTERS.get(ext)

    @classmethod
    def supports_format(cls, file_path: Path) -> bool:
        """Check if format is supported"""
        return file_path.suffix.lower() in cls.CONVERTERS

    @classmethod
    def convert_to_text(cls, file_path: Path) -> Dict[str, Any]:
        """
        Convert file to text

        Returns:
            Dict with 'text' and 'metadata'
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        converter = cls.get_converter(file_path)
        if not converter:
            raise ValueError(
                f"Unsupported file format: {file_path.suffix}\n"
                f"Supported formats: {', '.join(cls.CONVERTERS.keys())}"
            )

        try:
            text = converter.convert(file_path)
            return {
                "text": text,
                "metadata": {
                    "file": file_path.name,
                    "type": "converted",
                    "format": file_path.suffix.lower(),
                    "size": file_path.stat().st_size,
                    "converter": converter.__name__
                }
            }
        except Exception as e:
            raise RuntimeError(
                f"Error converting {file_path.name}: {str(e)}"
            ) from e

    @classmethod
    def list_supported_formats(cls) -> List[str]:
        """List all supported formats"""
        return sorted(cls.CONVERTERS.keys())


def convert_file(file_path: Path) -> Dict[str, Any]:
    """
    Convenience function to convert any supported file to text

    Args:
        file_path: Path to file

    Returns:
        Dict with 'text' and 'metadata'
    """
    return ConverterRegistry.convert_to_text(file_path)


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Format Converter - Extract text from documents")
        print("\nUsage:")
        print("  python3 format_converters.py <file>")
        print("\nSupported formats:")
        for fmt in ConverterRegistry.list_supported_formats():
            print(f"  {fmt}")
        sys.exit(1)

    file_path = Path(sys.argv[1])

    try:
        result = convert_file(file_path)
        print("="*80)
        print(f"File: {result['metadata']['file']}")
        print(f"Format: {result['metadata']['format']}")
        print(f"Converter: {result['metadata']['converter']}")
        print("="*80)
        print(result['text'])
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)
